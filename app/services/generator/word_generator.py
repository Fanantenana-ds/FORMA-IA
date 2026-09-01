# ============================================================
# GÉNÉRATEUR WORD — TDR (AVEC FORMATAGE PROFESSIONNEL)
# ============================================================

import os
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import logging

logger = logging.getLogger(__name__)


class WordGenerator:
    """Générateur de documents Word professionnels pour les TDR"""

    # Couleurs ALTIORA
    COLOR_PRIMARY = RGBColor(0, 51, 102)      # Bleu foncé
    COLOR_SECONDARY = RGBColor(0, 102, 204)   # Bleu clair
    COLOR_ACCENT = RGBColor(255, 165, 0)      # Orange
    COLOR_TEXT = RGBColor(0, 0, 0)            # Noir

    def __init__(self, output_dir: str = "exports/tdr"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        logger.info(f"📁 Stockage Word: {self.output_dir}")

    def generate(self, tdr_data: Dict[str, Any], brief: Dict[str, Any]) -> str:
        """
        Génère un document Word professionnel à partir des données du TDR

        Args:
            tdr_data: Données du TDR généré par Groq
            brief: Brief client original

        Returns:
            str: Chemin vers le fichier généré
        """
        try:
            doc = Document()

            # 1. En-tête avec logo
            self._add_header(doc)

            # 2. Titre principal
            title = tdr_data.get("tdr", {}).get("titre", "Termes de Référence")
            self._add_title(doc, title)

            # 3. Informations de base
            self._add_info_block(doc, brief)

            # 4. Sections
            sections = tdr_data.get("tdr", {}).get("sections", {})
            section_titles = {
                "contexte": "1. Contexte et justification",
                "objectifs": "2. Objectifs de la formation",
                "public": "3. Public cible et prérequis",
                "contenu": "4. Contenu et programme détaillé",
                "methodologie": "5. Méthodologie pédagogique",
                "evaluation": "6. Évaluation et suivi",
                "budget": "7. Budget et planning",
                "annexes": "8. Annexes"
            }

            for key, title_text in section_titles.items():
                content = sections.get(key, "À préciser par le client")
                self._add_section(doc, title_text, content)

            # 5. Pied de page
            self._add_footer(doc)

            # 6. Sauvegarder
            client_name = brief.get('client', 'client').replace(' ', '_')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"TDR_{client_name}_{timestamp}.docx"
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)

            logger.info(f"✅ Word généré: {filepath}")
            return filepath

        except Exception as e:
            logger.error(f"❌ Erreur génération Word: {e}")
            raise

    def _add_header(self, doc):
        """Ajoute l'en-tête avec logo et nom de l'entreprise"""
        # En-tête
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "ALTIORA SOLUTIONS"
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_para.runs[0]
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = self.COLOR_PRIMARY

        # Ligne de séparation
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()

    def _add_title(self, doc, title: str):
        """Ajoute le titre principal formaté"""
        # Titre
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = self.COLOR_PRIMARY

        # Sous-titre
        sub = doc.add_paragraph("Termes de Référence")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.runs[0]
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = self.COLOR_SECONDARY

        doc.add_paragraph()

    def _add_info_block(self, doc, brief: Dict[str, Any]):
        """Ajoute le bloc d'informations"""
        # Tableau d'informations
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Données du tableau
        info_data = [
            ("Client", brief.get('client', 'Non précisé')),
            ("Date", datetime.now().strftime('%d/%m/%Y')),
            ("Durée", brief.get('duree', 'Non précisé')),
            ("Format", brief.get('format', 'Non précisé'))
        ]

        for i, (label, value) in enumerate(info_data):
            # Label
            cell_label = table.cell(i, 0)
            cell_label.text = label
            cell_label.paragraphs[0].runs[0].bold = True
            cell_label.paragraphs[0].runs[0].font.color.rgb = self.COLOR_PRIMARY

            # Valeur
            cell_value = table.cell(i, 1)
            cell_value.text = value

        doc.add_paragraph()

    def _add_section(self, doc, title: str, content: str):
        """Ajoute une section du document avec formatage"""
        if not content or content == "À préciser par le client":
            content = "Information à préciser par le client."

        # Titre de section
        heading = doc.add_heading(title, level=1)
        run = heading.runs[0]
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = self.COLOR_SECONDARY

        # Traiter le contenu
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith('-') or line.startswith('•') or line.startswith('*'):
                # Liste à puces
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.5)
                p.add_run(line.lstrip('-•* '))
            elif line.startswith('Module') or line.startswith('1.') or line.startswith('2.') or line.startswith('3.'):
                # Titre de module
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(line)
                run.bold = True
                run.font.color.rgb = self.COLOR_PRIMARY
            elif line.startswith('Annexe'):
                # Titre d'annexe
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.color.rgb = self.COLOR_ACCENT
            else:
                # Paragraphe normal
                p = doc.add_paragraph(line)
                p.paragraph_format.first_line_indent = Inches(0.3)

        # Saut de ligne entre sections
        doc.add_paragraph()

    def _add_footer(self, doc):
        """Ajoute le pied de page"""
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"ALTIORA SOLUTIONS — Document confidentiel — {datetime.now().year}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.runs[0]
        run.font.size = Pt(8)
        run.font.color.rgb = self.COLOR_TEXT